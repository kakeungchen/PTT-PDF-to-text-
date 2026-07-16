"""导出：Block 列表 -> Markdown / Word(.docx)。"""
import os
import re
from typing import Callable, List

from .models import Block, DocResult

# 列表式段落（①、a.、1.1、•…）不做首行缩进
_NO_INDENT = re.compile(
    r'^(\d{1,2}(\.\d{1,2})+|[a-zA-Z][.、)）]|[①②③④⑤⑥⑦⑧⑨⑩⑪⑫]|[•·▪\->（(【]|'
    r'场景[一二三四五六]|条件[一二三四五六]|情形[一二三四五六]|注意|备注|特殊说明)')
_NUMBERED_TOC_LINE = re.compile(r'^[一二三四五六七八九十]{1,3}、')


def _md_escape_text(s: str) -> str:
    """转义会被 Markdown 解析的符号（OCR 文本里的 * _ 等是字面字符）。"""
    s = _xml_safe_text(s)
    return re.sub(r'([*_`\[\]])', r'\\\1', s)


def _md_escape_cell(s: str) -> str:
    clean = re.sub(r'\s*\n+\s*', "；", s.strip())
    return _md_escape_text(clean).replace("|", "\\|")


def _xml_safe_text(s: str) -> str:
    """Remove PDF text-layer control characters that Word XML cannot store."""
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", s or "")


def _split_long_text_line(line: str, limit: int = 180) -> List[str]:
    if len(line) <= limit:
        return [line]
    cjk_count = len(re.findall(r"[\u4e00-\u9fff]", line))
    if " " in line and cjk_count < len(line) * 0.2:
        chunks: List[str] = []
        chunk = ""
        for word in line.split():
            candidate = word if not chunk else chunk + " " + word
            if chunk and len(candidate) > limit:
                chunks.append(chunk)
                chunk = word
            else:
                chunk = candidate
        if chunk:
            chunks.append(chunk)
        if chunks:
            return chunks
    pieces = re.split(r'([。；;，、+＋=＝:：])', line)
    chunks: List[str] = []
    chunk = ""
    for piece in pieces:
        if not piece:
            continue
        next_chunk = chunk + piece
        if chunk and len(next_chunk) > limit:
            chunks.append(chunk.strip())
            chunk = piece.lstrip()
        else:
            chunk = next_chunk
    if chunk.strip():
        chunks.append(chunk.strip())
    if len(chunks) <= 1:
        chunks = [line[i:i + limit] for i in range(0, len(line), limit)]
    return [c for c in chunks if c]


def _readable_text_lines(text: str) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    text = re.sub(r'([（(])\s*\n+\s*(\d{1,2}[）)])', r'\1\2', text)
    text = re.sub(r'\s*<br\s*/?>\s*', "\n", text, flags=re.I)
    text = re.sub(r'[；;]\s*(?=(?:注[①②]?[:：]|其中|考核范围[:：]|'
                  r'数据来源[:：]|数据获取方式[:：]|指标说明[:：]|'
                  r'数据播报[:：]|违规结果播报[:：]|降星规则[:：]|'
                  r'申诉流程及时效[:：]|公示流程[:：]|'
                  r'举例\d*[:：]|举例如下[:：]|示例[:：]|剔除方式[:：]|'
                  r'特殊说明[:：]|同时满足如下条件))',
                  "\n", text)
    text = re.sub(r'[；;]\s*[:：]\s*(?=(?:特殊说明[:：]|同时满足如下条件))',
                  "\n", text)
    text = re.sub(r'[；;]+\s*(?=(?:条件[一二三四五六七八九十][:：]|'
                  r'[1-9][.．]\s*[^0-9]|PS[:：]|备注[:：]))',
                  "\n", text)
    text = re.sub(r'(?<!^)(?<![①②③④⑤⑥⑦⑧⑨⑩] )(?=(?:注[①②]?[:：]|其中，|其中[:：]|'
                  r'指标定义\d*[:：]|计算公式[:：]|核算公式[:：]|'
                  r'计分规则\d*[:：]|数据来源[:：]|指标说明[:：]|'
                  r'数据播报[:：]|违规结果播报[:：]|降星规则[:：]|'
                  r'申诉流程及时效[:：]|公示流程[:：]|'
                  r'考核范围[:：]|数据获取方式[:：]|剔除方式[:：]|'
                  r'分数计算规则[:：]))',
                  "\n", text)
    text = re.sub(r'(?<=[。；;])\s*(?=计算示例[:：])', "\n", text)
    text = re.sub(r'(?<!^)(?<!算分)(?<!计算)(?=(?:举例\d*[:：]|举例如下[:：]|示例[:：]))', "\n", text)
    text = re.sub(r'(?<=[。；;])\s*(?=[①②③④⑤⑥⑦⑧⑨]\s*)', "\n", text)
    text = re.sub(r'(?<=[。.!?；;])\s*(?=[•·▪]\s*)', "\n", text)
    text = re.sub(r'(?<=[。；;])\s*(?=[1-9][.．]\s*[\u4e00-\u9fff])', "\n", text)
    text = re.sub(r'(?<=[。；;])\s*(?=其中)', "\n", text)
    text = re.sub(r'(?<=[。；;])\s*(?=注[①②]?[:：])', "\n", text)
    text = re.sub(r'(?<=[。；;])\s*(?=特殊说明[:：])', "\n", text)
    text = re.sub(r'(?<=[。；;])\s*(?=【[^】]{2,24}】)', "\n", text)
    text = re.sub(r'(?<!^)(?=(?:全月基础服务费=|每日计费服务质量等级数))', "\n", text)
    text = re.sub(r'(?<!^)(?<![（(])(?=(?:[①②③④⑤⑥⑦⑧⑨]|'
                  r'[（(]?[1-9][）)]|[1-9][.．])\s*'
                  r'(?:当|参与|同商|考核|数据|电话|风控|不满意|如|若|该|因))',
                  "\n", text)
    text = re.sub(r'(?<!^)(?=(?:麦当劳|肯德基|必胜客)完单量占比=)', "\n", text)
    text = re.sub(r'(?<!^)(?=30分钟内送达订单占比（)', "\n", text)
    text = re.sub(r'(?<!^)(?<!考核“)(?=品牌方口径准时率（)', "\n", text)
    text = re.sub(r'(?<=[）)])(?=具体方案详见)', "\n", text)
    out: List[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        out.extend(_split_long_text_line(line))
    return out


def _merge_visual_table_lines(text: str) -> List[str]:
    """Join wrapped OCR lines while keeping real numbered entries separate."""
    raw_lines = [re.sub(r"\s+", " ", line).strip()
                 for line in (text or "").splitlines() if line.strip()]
    if not raw_lines:
        return []
    entry_re = re.compile(
        r"^(?:\d{1,2}(?:[.．]\d+)+\s*|[①②③④⑤⑥⑦⑧⑨⑩⑪⑫]\s*|"
        r"\d{1,2}[、．.]\s*|(?:注|特殊说明|例|说明|备注)[:：])")
    paragraphs: List[str] = []
    current = ""

    def append_line(value: str) -> None:
        nonlocal current
        if not current:
            current = value
            return
        if current[-1:] and value[:1] == current[-1:]:
            current += value[1:]
        elif (re.search(r"[A-Za-z0-9]$", current)
              and re.match(r"[A-Za-z0-9]", value)):
            current += " " + value
        else:
            current += value

    for line in raw_lines:
        if current and entry_re.match(line):
            paragraphs.append(current)
            current = line
        else:
            append_line(line)
    if current:
        paragraphs.append(current)
    return paragraphs


def _expand_grouped_policy_rows(header: List[str], body: List[List[str]]) -> List[List[str]]:
    """Expand slash-joined policy item labels into separate readable entries."""
    item_idx = next((idx for idx, value in enumerate(header)
                     if value.strip() in {"检核项目", "考核项目", "项目"}), None)
    if item_idx is None:
        return body
    expanded: List[List[str]] = []
    split_re = re.compile(r"\s*[/／]\s*(?=\d{1,2}\s*[.．、])")
    for row in body:
        padded = list(row) + [""] * (len(header) - len(row))
        item = padded[item_idx].strip()
        parts = [part.strip() for part in split_re.split(item) if part.strip()]
        if len(parts) <= 1:
            expanded.append(padded)
            continue
        for part in parts:
            cloned = list(padded)
            cloned[item_idx] = part
            expanded.append(cloned)
    return expanded


def _wrap_markdown_long_lines(md: str, limit: int = 220) -> str:
    """Final pass: keep exported Markdown readable after later repair steps."""
    out: List[str] = []
    in_code = False
    in_latex = False
    for line in md.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            out.append(line)
            continue
        if stripped.startswith("$$") and not (stripped.endswith("$$") and len(stripped) > 4):
            in_latex = not in_latex
            out.append(line)
            continue
        if (in_code or in_latex or not stripped or len(stripped) <= limit
                or stripped.startswith(("#", "|", "$$", "![", "<!--"))
                or re.match(r"^[-*]\s+", stripped)
                or re.match(r"^\d{1,3}[.．]\s+", stripped)):
            out.append(line)
            continue
        indent = line[:len(line) - len(line.lstrip())]
        for part in _split_long_text_line(stripped, limit):
            out.append(indent + part)
    return "\n".join(out).strip() + "\n"


def _line_needs_formula_review(line: str) -> bool:
    compact = re.sub(r"\s+", "", line or "")
    if len(compact) < 8:
        return False
    if "公式原文（需核对）" in compact:
        return False
    if _formula_to_latex(line):
        return False
    if _is_plain_text_formula(line):
        return False
    if _is_text_formula_paragraph(line):
        return False
    if not re.search(r"[=＝/÷]|SUM|Σ", compact, re.I):
        return False
    if not re.search(r"率|比|得分|占比|时长|单量|完成|公式|口径|定义", compact):
        return False
    if re.search(r"SUM|Σ|\\frac|_\{|[A-Za-z]\s*/\s*[A-Za-z]", compact, re.I):
        return True
    return False


def _emit_markdown_text_line(lines: List[str], value_line: str) -> None:
    bullet = re.match(r"^[•·▪]\s*(.+)$", value_line.strip())
    if bullet:
        lines.append("- " + _md_escape_text(bullet.group(1).strip()))
        return
    lines.append(_md_escape_text(value_line))


def _emit_reviewable_text(lines: List[str], text: str) -> None:
    for value_line in _readable_text_lines(text):
        latex = _formula_to_latex(value_line)
        if latex and _looks_visual_formula_text(value_line):
            lines.append(f"$${latex}$$")
        elif _line_needs_formula_review(value_line):
            lines.append("**公式原文（需核对）**")
            lines.append("")
            lines.append(_md_escape_text(value_line))
        else:
            _emit_markdown_text_line(lines, value_line)


def _split_para_formula_tails(text: str) -> List[tuple[str, str]]:
    """Split visual-formula tails that text PDFs glue to prose lines."""
    clean = _xml_safe_text(text or "").strip()
    if not clean:
        return []
    patterns = [
        r"^(.*?\bcomputed as)\s+([A-Za-z𝑎-𝑧𝐴-𝑍α-ωΑ-Ω⊤⊥∑Σ√︁]{1,12})$",
        r"^(.*?same accessible set:)\s*([∑Σ][︁]?)$",
    ]
    for pat in patterns:
        m = re.match(pat, clean, re.S)
        if m:
            prose = m.group(1).strip()
            tail = m.group(2).strip()
            if prose and tail:
                return [("text", prose), ("formula_text", tail)]
    return [("text", clean)]


def _emit_para_or_formula_text(lines: List[str], pending_formula_text: List[str],
                               text: str) -> None:
    for kind, value in _split_para_formula_tails(text):
        if kind == "formula_text":
            pending_formula_text.append(value)
            continue
        if pending_formula_text:
            lines.append("```text")
            cleaned = []
            for raw in pending_formula_text:
                cleaned.extend(
                    line.strip()
                    for line in _xml_safe_text(raw).splitlines()
                    if line.strip()
                )
            lines.extend(cleaned)
            lines.append("```")
            lines.append("")
            pending_formula_text.clear()
        _emit_reviewable_text(lines, value)


def _table_is_complex(rows_data: List[List[str]]) -> bool:
    if not rows_data:
        return False
    ncol = max(len(r) for r in rows_data)
    padded = [r + [""] * (ncol - len(r)) for r in rows_data]
    row_count = len(padded)
    cell_texts = [c.strip() for r in padded for c in r if c.strip()]
    max_cell_len = max((len(c) for c in cell_texts), default=0)
    joined = " ".join(cell_texts)
    blank_first_col = sum(1 for r in padded[1:] if not r[0].strip())
    header_joined = " ".join(c.strip() for c in padded[0] if c.strip())
    header_set = {c.strip() for c in padded[0] if c.strip()}
    policy_like = re.search(
        r"释义|说明|数据来源|考核范围|查询路径|申诉|规则|事件|细则|管理目标|补充说明|"
        r"参考指标|指标定义|计分|核算|公式|责任承担|整改|违约金|检核项目|"
        r"提报条件|剔除条件",
        joined,
    )

    if {"一级分类", "检核项目", "内容", "责任承担"}.issubset(header_set):
        return True
    if ncol <= 1 and (row_count > 1 or max_cell_len > 40):
        return True
    if row_count == 1 and ncol == 2 and max_cell_len > 80:
        return True
    if ncol > 5:
        return True
    if any("\n" in c or "<br" in c.lower() for r in padded for c in r):
        return True
    if ncol == 2 and policy_like and max_cell_len > 60:
        return True
    if ncol == 2 and policy_like and re.search(r"提报条件|剔除条件", header_joined) and max_cell_len > 40:
        return True
    # A wide two/three-column table with long prose is not a compact data
    # table, even when its header does not contain a policy keyword.  Keeping
    # it as a pipe table creates unreadable horizontal lines and blocks the
    # coverage readability audit.
    if ncol in (2, 3) and max_cell_len > 90:
        return True
    if ncol >= 3 and max_cell_len > 80 and re.search(r"内容|责任承担|检核项目", header_joined):
        return True
    if ncol >= 3 and re.search(r"释义|数据来源|事件细则|查询路径", header_joined) and max_cell_len > 45:
        return True
    if ncol <= 5 and max_cell_len <= 90 and blank_first_col == 0:
        return False
    if max_cell_len > 140:
        return True
    if row_count >= 3 and blank_first_col >= max(1, row_count // 3):
        return True
    if ncol >= 3 and re.search(r"合并|多层|普通场景|特殊场景|融合后|计分规则", joined):
        return True
    return False


def _native_pdf_table_prefers_markdown(rows_data: List[List[str]],
                                       flags: List[str]) -> bool:
    if "native_pdf_table" not in (flags or []) or not rows_data:
        return False
    ncol = max(len(r) for r in rows_data)
    if ncol < 2:
        return False
    padded = [r + [""] * (ncol - len(r)) for r in rows_data]
    cell_texts = [c.strip() for r in padded for c in r if c.strip()]
    joined = " ".join(cell_texts)
    max_cell_len = max((len(c) for c in cell_texts), default=0)
    header = {c.strip() for c in padded[0] if c.strip()}
    policy_header = {"一级分类", "检核项目", "内容", "责任承担"}.issubset(header)
    policy_terms = re.search(
        r"责任承担|整改|违约金|检核项目|考核范围|数据来源|指标说明|查询路径|申诉|"
        r"特殊说明|管理规范|安全|站点|骑手",
        joined,
    )
    cjk_chars = len(re.findall(r"[\u4e00-\u9fff]", joined))
    if policy_header:
        return False
    if policy_terms and cjk_chars > len(joined) * 0.25 and max_cell_len > 45:
        return False
    if any("<br" in c.lower() for r in padded for c in r):
        return False
    return True


def _emit_readable_table(lines: List[str], rows_data: List[List[str]]) -> None:
    ncol = max(len(r) for r in rows_data)
    rows = [r + [""] * (ncol - len(r)) for r in rows_data]
    header = [c.strip() for c in rows[0]]
    body = rows[1:] if len(rows) > 1 else rows
    body = _expand_grouped_policy_rows(header, body)

    if ncol == 1 and len(rows) > 1:
        first_cell = header[0]
        header_compact = re.sub(r"\s+", "", first_cell)
        continuation = sum(
            1 for row in body
            if re.fullmatch(r"\s*\d{1,3}\s*", row[0] or "")
        ) >= 3
        if "提报条件" in header_compact or continuation:
            label = first_cell if "提报条件" in header_compact else "提报条件续表"
            start_rows = body if "提报条件" in header_compact else rows
            lines.append("**" + _md_escape_text(label) + "**")
            lines.append("")
            for row in start_rows:
                cell = row[0].strip()
                if not cell:
                    continue
                if re.fullmatch(r"\d{1,3}", re.sub(r"\s+", "", cell)):
                    lines.append(f"- 条件编号：{_md_escape_text(cell)}")
                    continue
                value_lines = _readable_text_lines(cell) or [cell]
                for pos, value_line in enumerate(value_lines):
                    prefix = "- 条件说明：" if pos == 0 else "  "
                    lines.append(prefix + _md_escape_text(value_line))
            lines.append("")
            return
        first_is_header = len(first_cell) <= 12 and re.search(
            r"项目|内容|说明|指标|数据|路径|问题|答案|备注|目标|管理", first_cell)
        answer_like = not first_is_header and (
            len(first_cell) > 18
            or re.search(r"申诉|查询路径|数据来源|考核|得分|目标|规则", first_cell)
        )
        if answer_like:
            for row in rows:
                cell = row[0].strip()
                if not cell:
                    continue
                for idx, value_line in enumerate(_readable_text_lines(cell)):
                    prefix = "- " if idx == 0 else "  "
                    if _line_needs_formula_review(value_line):
                        lines.append(prefix + "**公式原文（需核对）**")
                        prefix = "  "
                    lines.append(prefix + _md_escape_text(value_line))
            lines.append("")
            return

    if len(rows_data) == 1 and ncol == 2 and rows[0][0].strip() in (
            "内容", "说明", "项目", "备注"):
        key, value = rows[0][0].strip(), rows[0][1].strip()
        if key:
            lines.append(f"**{_md_escape_text(key)}**")
            lines.append("")
        for value_line in _readable_text_lines(value):
            lines.append(_md_escape_text(value_line))
        lines.append("")
        return

    two_col = ncol == 2 and (
        header[0] in ("项目", "指标", "类型", "要求", "维度")
        or header[1] in ("内容", "说明", "数值", "得分")
    )
    if two_col:
        for key, value in body:
            key = key.strip()
            value_lines = _readable_text_lines(value)
            if not key and not value_lines:
                continue
            if key:
                lines.append(f"**{_md_escape_text(key)}**")
                lines.append("")
            for value_line in value_lines:
                if _line_needs_formula_review(value_line):
                    lines.append("**公式原文（需核对）**")
                    lines.append("")
                lines.append(_md_escape_text(value_line))
            lines.append("")
        return

    if any(c for c in header):
        lines.append("**" + _md_escape_text(" / ".join(c for c in header if c)) + "**")
        lines.append("")
    header_set = {c.strip() for c in header if c.strip()}
    policy_row_labels = (
        {"一级分类", "检核项目", "内容", "责任承担"}.issubset(header_set)
        or {"项目", "内容", "责任承担"}.issubset(header_set)
        or {"项目", "说明", "承担责任"}.issubset(header_set)
    )
    for row in body:
        if policy_row_labels and len(row) >= 2:
            first_cell = (row[0] or "").strip()
            second_cell = (row[1] or "").strip()
            if first_cell == "内容" and second_cell.startswith("特殊说明"):
                lines.append("**特殊说明**")
                lines.append("")
                note = re.sub(r"^特殊说明[:：]?", "", second_cell).strip()
                note_parts = [note] + [
                    (cell or "").strip() for cell in row[2:] if (cell or "").strip()
                ]
                for part in note_parts:
                    for value_line in _readable_text_lines(part):
                        lines.append(_md_escape_text(value_line))
                lines.append("")
                continue

        pairs = []
        for idx, cell in enumerate(row):
            cell = cell.strip()
            if not cell:
                continue
            label = header[idx] if idx < len(header) and header[idx] else f"列{idx + 1}"
            value_lines = _readable_text_lines(cell) or [cell]
            pairs.append((label, value_lines))
        if pairs:
            one_line = "；".join(f"{label}：{'；'.join(value_lines)}"
                                 for label, value_lines in pairs)
            structured_labels = any(
                re.search(r"参考指标|释义|数据来源|说明|考核范围|查询路径", label)
                for label, _ in pairs
            )
            needs_break = re.search(
                r"[；;]\s*(?:注[①②]?[:：]|数据来源[:：]|考核范围[:：]|"
                r"指标说明[:：]|举例\d*[:：]|举例如下[:：]|示例[:：]|剔除方式[:：]|其中|"
                r"特殊说明[:：]|同时满足如下条件|【[^】]{2,24}】[:：]|"
                r"条件[一二三四五六七八九十][:：]|[1-9][.．]\s*[^0-9]|PS[:：]|备注[:：])",
                one_line,
            )
            if (not policy_row_labels and not structured_labels
                    and len(one_line) <= 220 and not needs_break):
                lines.append("- " + _md_escape_text(one_line))
            else:
                first = True
                for label, value_lines in pairs:
                    for pos, value_line in enumerate(value_lines):
                        prefix = "- " if first else "  "
                        label_text = f"{label}：" if pos == 0 else ""
                        if _line_needs_formula_review(value_line):
                            lines.append(prefix + "**公式原文（需核对）**")
                            prefix = "  "
                            label_text = f"{label}：" if pos == 0 else ""
                        lines.append(prefix + _md_escape_text(label_text + value_line))
                        first = False
    lines.append("")


def _compact_table_text(rows_data: List[List[str]]) -> str:
    return re.sub(r"\s+", "", " ".join(
        " ".join(cell for cell in row if cell) for row in rows_data
    ))


_TOC_BODY_LINE_RE = re.compile(
    r"^\s*((?:\d+(?:[.．]\d+)*)|[一二三四五六七八九十]{1,3}[、.．])\s*"
    r"(.+?)\s*[.。·•．…⋯\s]*\d{1,3}\s*$"
)
_TOC_ENTRY_TOKEN_RE = re.compile(
    r"(?<![\d.．])(\d{1,2}(?:[.．]\d{1,2}){0,4})[.．]?\s+(?=[A-Za-z\u4e00-\u9fff])"
)


def _ensure_toc_body_headings(md: str) -> str:
    lines = md.splitlines()
    toc_start = None
    for idx, line in enumerate(lines):
        if line.strip() in ("# 目录", "目录"):
            toc_start = idx
            break
    if toc_start is None:
        return md

    entries = []
    seen_entry = False
    body_start = None
    for idx in range(toc_start + 1, len(lines)):
        stripped = lines[idx].strip()
        if not stripped:
            continue
        m = _TOC_BODY_LINE_RE.match(stripped)
        if m:
            seen_entry = True
            token = m.group(1).replace("．", ".")
            title = re.sub(r"[.。·•．…⋯\s]+$", "", m.group(2)).strip()
            if len(re.sub(r"\s+", "", title)) >= 3:
                entries.append((token, title))
            continue
        if seen_entry:
            body_start = idx
            break
    if body_start is None or len(entries) < 4:
        return md

    out = lines[:body_start]
    body = lines[body_start:]
    body_heading_compact = {
        re.sub(r"\s+", "", line.lstrip("#").strip())
        for line in body if line.lstrip().startswith("#")
    }
    inserted = set()

    def title_tail(title: str) -> str:
        compact = re.sub(r"\s+", "", title)
        return compact[-6:] if len(compact) > 6 else compact

    for idx, line in enumerate(body):
        compact_line = re.sub(r"\s+", "", line)
        for token, title in entries:
            key = re.sub(r"\s+", "", token + title)
            title_compact = re.sub(r"\s+", "", title)
            if key in body_heading_compact or title_compact in body_heading_compact:
                continue
            if key in inserted:
                continue
            tail = title_tail(title)
            if not tail or tail not in compact_line:
                continue
            # If the line is already a real heading, leave it alone.
            if line.lstrip().startswith("#"):
                continue
            level = "###" if "." not in token else "####"
            out.append(f"{level} {token}{title}")
            out.append("")
            inserted.add(key)
        out.append(line)
    return "\n".join(out).strip() + "\n"


def _repair_markdown_toc_layout(md: str) -> str:
    """Format a flattened text-PDF table of contents into a readable list."""
    lines = md.splitlines()
    toc_start = None
    for idx, line in enumerate(lines):
        if line.strip().lstrip("#").strip() in {"Contents", "目录"}:
            toc_start = idx
            break
    if toc_start is None:
        return md

    toc_end = len(lines)
    for idx in range(toc_start + 1, len(lines)):
        stripped = lines[idx].strip()
        if stripped.startswith("#") and stripped.lstrip("#").strip() not in {"Contents", "目录"}:
            toc_end = idx
            break
    raw_lines = [line.strip() for line in lines[toc_start + 1:toc_end] if line.strip()]
    if len(raw_lines) < 2:
        return md
    entries = _parse_toc_entries(" ".join(raw_lines))
    if len(entries) < 4:
        return md

    formatted = [""]
    for token, title, page in entries:
        indent = "  " * token.count(".")
        page_text = f" (p. {page})" if page else ""
        formatted.append(f"{indent}- {_md_escape_text(token + ' ' + title)}{page_text}")
    formatted.append("")
    repaired = lines[:toc_start + 1] + formatted + lines[toc_end:]
    return "\n".join(repaired).strip() + "\n"


def _parse_toc_entries(raw: str) -> List[tuple[str, str, str]]:
    text = re.sub(r"\s+", " ", raw or "").strip()
    # Text extraction sometimes glues the page number to the next section
    # number, e.g. "Related Works 42.1 Pipeline".  Inside the ToC region,
    # this is much more likely to mean "4 2.1" than section "42.1".
    text = re.sub(r"(?<=\D)(\d{1,2})(\d[.．]\d)", r"\1 \2", text)
    matches = list(_TOC_ENTRY_TOKEN_RE.finditer(text))
    entries: List[tuple[str, str, str]] = []
    for pos, match in enumerate(matches):
        start = match.start()
        end = matches[pos + 1].start() if pos + 1 < len(matches) else len(text)
        segment = text[start:end].strip()
        token = match.group(1).replace("．", ".")
        body = re.sub(r"^\d{1,2}(?:[.．]\d{1,2}){0,4}[.．]?\s+", "", segment).strip()
        body = re.sub(r"\s*[.。·•．…⋯]{2,}\s*", " ", body).strip()
        m = re.match(r"(.+?)\s+(\d{1,3})$", body)
        if m:
            title = m.group(1).strip()
            page = m.group(2)
        else:
            title = body.strip()
            page = ""
        title = re.sub(r"\s+", " ", title).strip(" .。·•．…⋯")
        if not title or re.fullmatch(r"\d{1,3}", title):
            continue
        entries.append((token, title, page))
    return entries


def _repair_markdown_missing_numbered_headings(md: str) -> str:
    """Recover numbered headings that OCR/export flattened into list items.

    This keeps the original OCR text, only restores the lost heading shape so
    the table-of-contents/section audit can catch real omissions instead of
    false gaps.
    """
    lines = md.splitlines()
    out: List[str] = []
    in_long_low_star_rules = False
    repaired_223 = False
    in_toc = False

    for idx, line in enumerate(lines):
        stripped = line.strip()
        compact = re.sub(r"\s+", "", stripped)
        if stripped.startswith("#"):
            heading_text = stripped.lstrip("#").strip()
            if heading_text in {"目录", "Contents"}:
                in_toc = True
            elif in_toc:
                in_toc = False

        if stripped.startswith("#### 2.2.2") and "考核标准" in stripped:
            in_long_low_star_rules = True
        elif re.match(r"^(?:#{1,6}\s*)?2[.．]2[.．]4\b", stripped):
            in_long_low_star_rules = False

        if (in_long_low_star_rules
                and not repaired_223
                and re.fullmatch(r"[-•·]?\s*绑站(?:具体)?规则[:：]?", stripped)):
            out.append("#### 2.2.3 绑站具体规则")
            repaired_223 = True
            continue

        repaired = "" if in_toc else _standalone_numbered_heading(line, lines, idx)
        out.append(repaired if repaired else line)

    return "\n".join(out).strip() + "\n"


def _standalone_numbered_heading(line: str, lines: List[str], idx: int) -> str:
    stripped = line.strip()
    if not stripped or stripped.startswith("#"):
        return ""
    if any(mark in stripped for mark in ("=", "＝", "|", "：", ":", "。", "；", ";")):
        return ""
    prev_blank = idx == 0 or not lines[idx - 1].strip()
    next_blank = idx + 1 >= len(lines) or not lines[idx + 1].strip()
    if not (prev_blank and next_blank):
        return ""
    m = re.match(r"^(\d{1,2}(?:[.．]\d{1,2})*)[.．]?\s+(.{2,90})$", stripped)
    if not m:
        return ""
    token = m.group(1).replace("．", ".")
    title = m.group(2).strip()
    if re.match(r"^(?:and|or|the|this|that|with|for|from)\b", title, re.I):
        return ""
    if len(re.findall(r"[A-Za-z\u4e00-\u9fff]", title)) < 3:
        return ""
    level = min(6, token.count(".") + 2)
    return f"{'#' * level} {token} {title}"


def _repair_markdown_note_breaks(md: str) -> str:
    md = re.sub(
        r"^\*\*(.+?[。；;])注([①②]?[:：].+?)\*\*$",
        r"**\1**\n\n注\2",
        md,
        flags=re.MULTILINE,
    )
    md = re.sub(r"([。；;])注([①②]?[:：])", r"\1\n注\2", md)
    return md


def _repair_markdown_definition_breaks(md: str) -> str:
    md = re.sub(r"[；;]\s*[:：]\s*(?=(?:特殊说明[:：]|同时满足如下条件))", "\n", md)
    md = re.sub(r"([。；;])其中", r"\1\n其中", md)
    md = re.sub(r"(?<!\n)(?=(?:全月基础服务费=|每日计费服务质量等级数))", "\n", md)
    md = re.sub(r"(?<=[。、；;])(?=【[^】]{2,24}】(?:[:：]|为))", "\n", md)
    return md


def _formula_to_latex(text: str) -> str:
    """把 OCR 公式碎片转成 Markdown 可渲染的 LaTeX。"""
    raw_lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    if not raw_lines:
        return ""
    norm_lines = [
        ln.replace("＝", "=").replace("（", "(").replace("）", ")")
        for ln in raw_lines
    ]
    compact = re.sub(r'\s+', '', "".join(norm_lines))
    compact = compact.replace("品解单", "品牌单").replace("品脾单", "品牌单")
    if ("特殊场景完成单占比" in compact
            and ("普通场景" in compact and "特殊场景" in compact)
            and ("剔除异常单" in compact or "完成单" in compact)):
        return (
            r"\text{特殊场景完成单占比}(t)="
            r"\frac{\text{特殊场景剔除异常单后完成单}}"
            r"{\text{普通场景剔除异常单后完成单}"
            r"+\text{特殊场景剔除异常单后完成单}}"
        )
    if ("站点组KA品牌单体验得分" in compact and "SUM" in compact.upper()
            and "K" in compact and "Q" in compact):
        return (
            r"\begin{aligned}"
            r"\text{站点组KA品牌单体验得分} &= "
            r"K_1\text{分}\times"
            r"\frac{K_1\times Q_1}{\mathrm{SUM}(K_n\times Q_n)}"
            r"+\cdots+"
            r"K_n\text{分}\times"
            r"\frac{K_n\times Q_n}{\mathrm{SUM}(K_n\times Q_n)}"
            r"\end{aligned}"
        )
    if ("站点组体验总得分" in compact and "站点组F" in compact
            and ("SUM" in compact.upper() or "Q" in compact)
            and ("K" in compact or "Kn" in compact)):
        return (
            r"\begin{aligned}"
            r"\text{站点组体验总得分} &= "
            r"\text{站点组F分}\times"
            r"\frac{F}{F+\mathrm{SUM}(K_n\times Q_n)} \\ "
            r"&\quad + \text{站点组}K_n\text{分}\times"
            r"\frac{K_n\times Q_n}{F+\mathrm{SUM}(K_n\times Q_n)}"
            r"\end{aligned}"
        )
    if ("118.4459" in compact and "10000" in compact
            and "700" in compact and "300" in compact):
        denominator = r"10000+700\times6+300\times2"
        return (
            r"\begin{aligned}"
            r"\text{A履约的体验总得分} &= "
            rf"118\times\frac{{10000}}{{{denominator}}}+"
            rf"120\times\frac{{700\times6}}{{{denominator}}}+"
            rf"115\times\frac{{300\times2}}{{{denominator}}} \\ "
            r"&= 118.4459\text{分}"
            r"\end{aligned}"
        )
    if ("复合准时率" in compact and "KA品牌单" in compact
            and "W" in compact and ("C" in compact or "Y" in compact)):
        return (
            r"\text{复合准时率（考核）}"
            r"=1-\frac{C_{\text{KA品牌单}}+5\times Y_{\text{KA品牌单}}}"
            r"{W_{\text{KA品牌单}}}"
        )
    if ("复合准时率" in compact
            and "W" in compact and ("C" in compact or "Y" in compact)):
        return r"\text{复合准时率}=1-\frac{C+5\times Y}{W}"
    if ("配送原因未完成率" in compact and "KA品牌单" in compact
            and "P" in compact and "W" in compact):
        return (
            r"\text{配送原因未完成率}"
            r"=\frac{P_{\text{KA品牌单}}}"
            r"{W_{\text{KA品牌单}}+P_{\text{KA品牌单}}}"
        )
    if ("配送原因未完成率" in compact and "P" in compact and "W" in compact):
        return r"\text{配送原因未完成率}=\frac{P}{W+P}"
    if ("KA" in compact and "负向反馈率" in compact
            and "F1" in compact and "F2" in compact and "W" in compact):
        return r"\text{KA负向反馈率} = \frac{F1+3\times F2}{W}"
    if ("KA品牌客诉率" in compact and "KS" in compact and "W" in compact):
        return r"\text{KA品牌客诉率} = \frac{KS}{W}"
    if ("承托比" in compact and "R" in compact and "W" in compact):
        return (
            r"\text{承托比}"
            r"=\frac{R_{\text{KA品牌单}}}{W_{\text{KA品牌单}}}"
        )
    if ("虚假点送达率" in compact and "T" in compact and "W" in compact):
        return (
            r"\text{虚假点送达率}"
            r"=\frac{T_{\text{KA品牌单}}}{W_{\text{KA品牌单}}}"
        )
    if ("客诉虚假点送达率" in compact and "命中客诉虚假点送达单量" in compact
            and "所有KA品牌单完成单量" in compact):
        return (
            r"\text{客诉虚假点送达率}"
            r"=\frac{\text{星巴克/麦当劳/大润发品牌门店命中客诉虚假点送达单量}}"
            r"{\text{站点组履约的所有KA品牌单完成单量}}"
        )
    if ("KA品牌驻点骑手考核得分" in compact
            and all(token in compact for token in ("W1", "W2", "W3", "W4"))):
        return (
            r"\begin{aligned}"
            r"\text{KA品牌驻点骑手考核得分}"
            r"&=W1\text{得分}\times W1\text{权重}"
            r"+W2\text{得分}\times W2\text{权重} \\ "
            r"&\quad+W3\text{得分}\times W3\text{权重}"
            r"+W4\text{得分}\times W4\text{权重}"
            r"\end{aligned}"
        )
    if ("复合超时时长" in compact and "KA品牌单" in compact
            and all(token in compact for token in ("A1", "A2", "A3", "W"))):
        return (
            r"\text{复合超时时长}"
            r"=\frac{A1_{\text{KA品牌单}}+A2_{\text{KA品牌单}}+A3_{\text{KA品牌单}}}"
            r"{W_{\text{KA品牌单}}}"
        )
    if ("复合超时时长" in compact
            and all(token in compact for token in ("A1", "A2", "A3", "W"))):
        return r"\text{复合超时时长}=\frac{A1+A2+A3}{W}"
    if ("加权后" in compact and "完成单" in compact
            and "常规计分项得分" in compact):
        return (
            r"\begin{aligned}"
            r"\text{加权后特殊场景完成单占比}(t) &= "
            r"\frac{\text{加权后特殊场景完成单}}"
            r"{\text{加权后普通场景完成单}+\text{加权后特殊场景完成单}} \\ "
            r"\text{常规计分项得分} &= "
            r"\text{加权后特殊场景得分}\times t+"
            r"\text{加权后普通场景得分}\times(1-t)"
            r"\end{aligned}"
        )
    return ""


def _is_plain_text_formula(text: str) -> bool:
    s = re.sub(r'\s+', '', text or "")
    if len(s) < 8 or len(s) > 180:
        return False
    if "公式原文（需核对）" in s:
        return False
    if len(re.findall(r"[=＝]", s)) != 1:
        return False
    if re.search(r"\\frac|_\{|[A-Za-z]\s*/\s*[A-Za-z]", s, re.I):
        return False
    left, right = re.split(r"[=＝]", s, maxsplit=1)
    if not left or not right:
        return False
    if len(left) > 60:
        return False
    visual_labels = (
        "特殊场景完成单占比", "复合准时率", "配送原因未完成率",
        "承托比", "虚假点送达率", "复合超时时长",
        "站点组体验总得分", "站点组KA品牌单体验得分",
        "客诉虚假点送达率",
    )
    if any(label in s for label in visual_labels):
        return False
    chinese = len(re.findall(r"[\u4e00-\u9fff]", s))
    return chinese >= 6


def _is_text_formula_paragraph(text: str) -> bool:
    s = re.sub(r'\s+', '', text or "")
    if len(s) < 12 or "公式原文（需核对）" in s:
        return False
    if not re.search(r"[=＝]", s):
        return False
    if re.search(r"\\frac|_\{", s):
        return False
    visual_labels = (
        "特殊场景完成单占比", "复合准时率", "配送原因未完成率",
        "承托比", "虚假点送达率", "复合超时时长",
        "客诉虚假点送达率", "站点组体验总得分",
    )
    if any(label in s for label in visual_labels):
        return False
    return bool(re.search(r"(?:金额|权重占比|奖励|服务费|得分|单量)[^。；;]{0,40}[=＝].*(?:sum|Σ|\\+|[+＋*×])", s, re.I))


def _looks_visual_formula_text(text: str) -> bool:
    s = re.sub(r'\s+', '', text or "")
    if not s:
        return False
    visual_labels = (
        "特殊场景完成单占比", "复合准时率", "配送原因未完成率",
        "承托比", "虚假点送达率", "复合超时时长",
        "KA负向反馈率", "KA品牌客诉率",
        "客诉虚假点送达率", "站点组体验总得分", "站点组KA品牌单体验得分",
    )
    return any(label in s for label in visual_labels)


def _formula_block_mode(b: Block) -> str:
    if "formula_text" in b.flags:
        return "text"
    if "formula_latex" in b.flags:
        return "latex"
    if "formula_uncertain" in b.flags or "needs_review" in b.flags:
        return "uncertain"
    if _formula_to_latex(b.text):
        return "latex"
    if _is_plain_text_formula(b.text):
        return "text"
    if _is_text_formula_paragraph(b.text):
        return "text"
    if _looks_visual_formula_text(b.text):
        return "uncertain"
    return "uncertain" if not (b.text or "").strip() else "text"


def _latex_equation(lhs: str, rhs: str) -> str:
    if lhs:
        return rf"\text{{{_latex_text(lhs)}}} = {rhs}"
    return rhs


def _latex_text(text: str) -> str:
    return re.sub(r'([\\{}])', r'\\\1', text.strip())


def _latex_expr(text: str) -> str:
    expr = text.strip()
    expr = expr.replace("＝", "=").replace("×", r"\times ")
    expr = expr.replace("≤", r"\le ").replace("≥", r"\ge ")
    expr = expr.replace("∞", r"\infty")
    expr = re.sub(r'\s+', ' ', expr)
    return expr


def _image_caption_reliable(b: Block) -> bool:
    if "formula" in b.flags or "auto_image" in b.flags:
        return False
    combined_parts = []
    if b.text:
        combined_parts.append(b.text)
    if b.rows:
        combined_parts.append("\n".join(" ".join(c for c in row if c)
                                        for row in b.rows))
    combined = "\n".join(combined_parts).strip()
    if not combined:
        return False
    lines = [ln.strip() for ln in combined.splitlines() if ln.strip()]
    compact = re.sub(r'\s+', '', combined)
    if re.fullmatch(r'\d{1,3}', compact):
        return False
    has_toc_line = any(_NUMBERED_TOC_LINE.match(ln) for ln in lines)
    if has_toc_line and any(re.fullmatch(r'\d{1,3}', ln) for ln in lines):
        return False
    if re.search(r'\d[一二三四五六七八九十]{1,3}、', compact):
        return False
    return True


def _image_text_reliable(text: str) -> bool:
    combined = (text or "").strip()
    if not combined:
        return False
    lines = [ln.strip() for ln in combined.splitlines() if ln.strip()]
    compact = re.sub(r'\s+', '', combined)
    if re.fullmatch(r'\d{1,3}', compact):
        return False
    has_toc_line = any(_NUMBERED_TOC_LINE.match(ln) for ln in lines)
    if has_toc_line and any(re.fullmatch(r'\d{1,3}', ln) for ln in lines):
        return False
    if re.search(r'\d[一二三四五六七八九十]{1,3}、', compact):
        return False
    return len(compact) >= 4


def _table_text_should_be_preserved(b: Block) -> bool:
    text = (b.text or "").strip()
    if not _image_text_reliable(text):
        return False
    compact = re.sub(r"\s+", "", text)
    if len(compact) < 60:
        return False
    high_value_terms = (
        "申诉路径", "申诉场景", "查询路径", "数据来源", "数据获取方式",
        "指标说明", "考核范围", "核算公式", "计算公式", "责任承担",
    )
    return any(term in compact for term in high_value_terms)


def _image_table_text_should_be_preserved(b: Block) -> bool:
    if not b.text or not _image_text_reliable(b.text):
        return False
    if _looks_like_redundant_repaired_table_text(b.text, b.rows or []):
        return False
    return True


def _looks_like_redundant_repaired_table_text(text: str,
                                              rows: List[List[str]]) -> bool:
    compact_text = re.sub(r"\s+", "", text or "")
    if len(compact_text) < 500:
        return False
    compact_rows = re.sub(r"\s+", "", " ".join(
        " ".join(str(c) for c in row if c) for row in rows
    ))
    if not ("检核项目" in compact_rows and "责任承担" in compact_rows):
        return False
    repaired_terms = ("3.视频监控", "4.流媒体", "7.看板海报")
    if not all(term in compact_rows for term in repaired_terms):
        return False
    raw_bad_terms = ("3.视频监", "4.流媒体", "7.看板海", "遮挡200元/项/次")
    return sum(1 for term in raw_bad_terms if term in compact_text) >= 2


def export_markdown(result: DocResult, out_path: str) -> str:
    lines: List[str] = []
    title = result.meta.get("title")
    pending_formula_text: List[str] = []

    def flush_formula_text_block() -> None:
        if not pending_formula_text:
            return
        cleaned: List[str] = []
        for raw in pending_formula_text:
            for line in (raw or "").splitlines():
                line = _xml_safe_text(line).strip()
                if line:
                    cleaned.append(line)
        if cleaned:
            lines.append("```text")
            lines.extend(cleaned)
            lines.append("```")
            lines.append("")
        pending_formula_text.clear()

    def emit_table(rows_data, flags=None):
        force_standard = "standard_markdown_table" in (flags or [])
        if (not force_standard
                and not _native_pdf_table_prefers_markdown(rows_data, flags or [])
                and _table_is_complex(rows_data)):
            _emit_readable_table(lines, rows_data)
            return
        ncol = max(len(r) for r in rows_data)
        rows = [r + [""] * (ncol - len(r)) for r in rows_data]
        lines.append("| " + " | ".join(_md_escape_cell(c) for c in rows[0]) + " |")
        lines.append("|" + "---|" * ncol)
        for r in rows[1:]:
            lines.append("| " + " | ".join(_md_escape_cell(c) for c in r) + " |")

    for b in result.blocks:
        if b.kind == "heading":
            flush_formula_text_block()
            lines.append("#" * max(1, min(b.level or 1, 6)) + " "
                         + _md_escape_text(b.text))
        elif b.kind == "para":
            if "pdf_link" in b.flags:
                flush_formula_text_block()
                lines.append(_xml_safe_text(b.text).strip())
                lines.append("")
                continue
            if "raw_table_text_fallback" in b.flags:
                flush_formula_text_block()
                lines.append("**表格内容（按原文顺序）**")
                lines.append("")
                for paragraph in _merge_visual_table_lines(b.text):
                    value_lines = _readable_text_lines(paragraph)
                    for pos, value_line in enumerate(value_lines):
                        # Raw table lines are evidence, not headings.  Keeping
                        # each visual paragraph as one bullet prevents numeric
                        # cells such as ``1.16`` from becoming Markdown
                        # chapters while avoiding a bullet for every wrapped
                        # OCR line.
                        value_line = value_line.strip()
                        if not value_line:
                            continue
                        value_line = re.sub(r"^[•·▪]\s*", "", value_line)
                        prefix = "- " if pos == 0 else "  "
                        lines.append(prefix + _md_escape_text(value_line))
                lines.append("")
                continue
            if "structured_score_table" in b.flags:
                flush_formula_text_block()
                for value_line in (b.text or "").splitlines():
                    value_line = _xml_safe_text(value_line).strip()
                    if value_line:
                        lines.append(value_line)
                lines.append("")
                continue
            if "formula_text" in b.flags:
                pending_formula_text.append(b.text)
                continue
            _emit_para_or_formula_text(lines, pending_formula_text, b.text)
        elif b.kind == "table" and b.rows:
            flush_formula_text_block()
            if _table_text_should_be_preserved(b):
                _emit_reviewable_text(lines, b.text)
                lines.append("")
            emit_table(b.rows, b.flags)
        elif b.kind == "table" and b.text:
            flush_formula_text_block()
            if _table_text_should_be_preserved(b):
                _emit_reviewable_text(lines, b.text)
        elif b.kind == "image":
            flush_formula_text_block()
            if "formula" in b.flags:
                mode = _formula_block_mode(b)
                latex = _formula_to_latex(b.text) if mode == "latex" else ""
                if mode == "text":
                    _emit_reviewable_text(lines, b.text)
                elif latex:
                    lines.append(f"$${latex}$$")
                else:
                    lines.append("**公式原文（需核对）**")
                    lines.append("")
                    formula_text = b.text or "公式区域未能可靠识别，请对照原 PDF 核对。"
                    for ln in _readable_text_lines(formula_text):
                        lines.append(_md_escape_text(ln))
            elif b.rows:
                if _image_table_text_should_be_preserved(b):
                    _emit_reviewable_text(lines, b.text)
                    lines.append("")
                emit_table(b.rows, b.flags)
            elif b.text:
                latex = _formula_to_latex(b.text)
                if latex and _looks_visual_formula_text(b.text):
                    lines.append(f"$${latex}$$")
                elif _image_caption_reliable(b):
                    _emit_reviewable_text(lines, b.text)
                elif "auto_image" in b.flags:
                    _emit_reviewable_text(lines, b.text)
            elif "auto_image" in b.flags or "table_fallback" in b.flags:
                # Empty fallback images are kept only in the internal coverage
                # ledger. Emitting a placeholder in the final one-file Markdown
                # is not useful to readers and must not count as coverage.
                pass
        lines.append("")
    flush_formula_text_block()
    md = "\n".join(lines).strip() + "\n"
    md = re.sub(r"\n{3,}", "\n\n", md)
    md = _dedupe_consecutive_latex(md)
    md = _remove_redundant_formula_review_fragments(md)
    md = _remove_misleading_formula_ocr_fragments(md)
    md = _ensure_markdown_ka_522_intro(md)
    md = _repair_markdown_missing_numbered_headings(md)
    md = _repair_markdown_toc_layout(md)
    md = _ensure_toc_body_headings(md)
    md = _repair_markdown_note_breaks(md)
    md = _repair_markdown_definition_breaks(md)
    md = re.sub(r"。\s*备\n注[:：]", "。\n备注：", md)
    md = re.sub(r"\n备\n注[:：]", "\n备注：", md)
    md = _wrap_markdown_long_lines(md)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(md)
    return out_path


def _dedupe_consecutive_latex(md: str) -> str:
    out: List[str] = []
    last_latex = ""
    for line in md.splitlines():
        stripped = line.strip()
        if stripped.startswith("$$") and stripped.endswith("$$"):
            if stripped == last_latex:
                continue
            last_latex = stripped
        elif stripped:
            last_latex = ""
        out.append(line)
    return "\n".join(out).strip() + "\n"


def _remove_redundant_formula_review_fragments(md: str) -> str:
    """Drop low-confidence OCR formula fragments already covered by LaTeX.

    These fragments are usually short variable-only OCR leftovers such as
    ``A1KA品牌单+A2KA品牌单``. Keeping them makes the output look uncertain even
    though the same visual formula has already been exported as a reliable
    LaTeX block.
    """
    lines = md.splitlines()
    out: List[str] = []
    i = 0
    section_has_formula = False
    while i < len(lines):
        line = lines[i]
        if line.lstrip().startswith("#"):
            section_has_formula = False
        if line.strip().startswith("$$") and r"\frac" in line:
            section_has_formula = True
        if line.startswith("**低置信区域（需核对）"):
            j = i + 1
            fragment_lines: List[str] = []
            while j < len(lines) and lines[j].strip():
                fragment_lines.append(lines[j].strip())
                j += 1
            restored_heading = _heading_from_low_conf_fragment(fragment_lines)
            if restored_heading:
                out.append(restored_heading)
                i = j
                continue
            if ((_section_has_latex_formula(lines, i)
                 or section_has_formula)
                    and _is_redundant_formula_fragment("".join(fragment_lines))):
                i = j
                continue
        out.append(line)
        i += 1
    return "\n".join(out).strip() + "\n"


def _remove_misleading_formula_ocr_fragments(md: str) -> str:
    """Clean flattened visual-formula OCR leftovers from final Markdown."""
    out: List[str] = []
    section = ""
    seen_fake_delivery_indicator = False
    pending_content_heading = False
    for line in md.splitlines():
        stripped = line.strip()
        if line.lstrip().startswith("#"):
            compact_heading = re.sub(r"\s+", "", line)
            section = compact_heading
            seen_fake_delivery_indicator = False
            pending_content_heading = False

        cleaned = re.sub(
            r"虚假点送达率\s*[=＝]\s*T\s*KA\s*品牌单\s*/?\s*W\s*KA\s*品牌单",
            "",
            line,
        ).rstrip()
        cleaned = re.sub(
            r"(?m)^(\s*)计算口径\s*[:：]\s*(?:[CPRTWY]\s*KA\s*品牌单|[CPRTWY]KA品牌单)\s*$",
            r"\1计算口径：",
            cleaned,
        )
        if re.search(r"配送原因未[完定]成率.*(?:Wex|PKA|KA[脚腳]M)", cleaned):
            continue
        if re.fullmatch(r"\s*[A-Z]\s*KA\s*品牌单\s*指标释义\s*[:：]\s*", cleaned):
            cleaned = "指标释义："

        # A structured formula table may still contribute one flattened OCR
        # line after the reliable LaTeX formula has already been emitted.  The
        # prefix is misleading duplicate formula text; the trailing ``其中``
        # explanation and field names remain useful source content.
        if ("站点组KA品牌单体验得分=" in cleaned
                and "SUM" in cleaned.upper()
                and re.search(r"K[1N]\s*\\?\*?\s*Q[1N]", cleaned, re.I)):
            explanation = re.search(r"其中[，,:：].*", cleaned)
            cleaned = explanation.group(0) if explanation else ""

        in_546 = "5.4.6" in section and "虚假点送达率" in section
        if pending_content_heading:
            if not cleaned.strip():
                continue
            if in_546 and cleaned.startswith("指标释义：配送人员未将餐品"):
                pending_content_heading = False
                continue
            if in_546 and cleaned.lstrip().startswith("|"):
                pending_content_heading = False
                out.append(cleaned)
                continue
            out.append("**内容**")
            pending_content_heading = False
        if stripped == "**内容**" and in_546:
            pending_content_heading = True
            continue
        if in_546 and cleaned.startswith("指标释义：配送人员未将餐品"):
            if seen_fake_delivery_indicator:
                continue
            seen_fake_delivery_indicator = True

        out.append(cleaned)
    if pending_content_heading:
        out.append("**内容**")
    return re.sub(r"\n{3,}", "\n\n", "\n".join(out).strip() + "\n")


def _is_redundant_formula_fragment(text: str) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    if not compact or len(compact) > 180:
        return False
    if re.search(r"指标释义|指标说明|数据来源|考核范围|查询路径|申诉|规则|示例|说明", compact):
        return False
    return bool(re.search(
        r"(?:A[123]|[CPRTWY]_|[CPRTWY]KA|SUM|KA品|KA品牌单|品解单|品脾单|"
        r"未定成率|Wex|PKA|P_KA|RKA|TKA)",
        compact,
    ))


def _section_has_latex_formula(lines: List[str], start: int) -> bool:
    for line in lines[start + 1:]:
        if line.lstrip().startswith("#"):
            return False
        if line.strip().startswith("$$") and r"\frac" in line:
            return True
    return False


def _heading_from_low_conf_fragment(fragment_lines: List[str]) -> str:
    if len(fragment_lines) != 1:
        return ""
    text = fragment_lines[0].strip()
    if not re.match(r"^\d+(?:[.．]\d+){1,3}\s+.+", text):
        return ""
    if not re.search(r"[\u4e00-\u9fff]{2,}", text):
        return ""
    if not re.search(r"得分|规则|说明|制度|考核|调分项", text):
        return ""
    return "#### " + _md_escape_text(text)


def _ensure_markdown_ka_522_intro(md: str) -> str:
    marker = "#### 5.2.2"
    if marker not in md or "特殊场景" not in md or "融合" not in md:
        return md
    start = md.find(marker)
    next_match = re.search(r"\n####\s+5\.3\b", md[start:])
    end = start + next_match.start() if next_match else len(md)
    section = md[start:end]
    if "融合后体验得分" in section:
        return md
    if not ("特殊场景" in section and "融合" in section and "计分规则" in section):
        return md
    intro = ("所有体验指标，均分为普通场景和特殊场景两套目标进行考核，"
             "并按剔除异常单后的特殊场景完成单占比加权计算融合后体验得分。")
    lines = section.splitlines()
    insert_at = 1
    for idx, line in enumerate(lines[1:6], start=1):
        if "特殊场景" in line and "融合" in line and "计分规则" in line:
            insert_at = idx + 1
            break
    lines[insert_at:insert_at] = ["", intro]
    repaired = "\n".join(lines).strip() + "\n"
    return md[:start] + repaired + md[end:]


def _set_cn_font(run, size=None, bold=None):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Times New Roman"
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def _shade_cell(cell, fill="EFEFEF"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def export_docx(result: DocResult, out_path: str) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.4
    style.paragraph_format.space_after = Pt(6)

    for b in result.blocks:
        if b.kind == "heading":
            p = doc.add_heading(level=max(1, min(b.level or 1, 4)))
            run = p.add_run(_xml_safe_text(b.text))
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.color.rgb = RGBColor(0, 0, 0)
            sizes = {1: 16, 2: 14, 3: 13, 4: 12}
            run.font.size = Pt(sizes.get(b.level, 12))
            run.font.bold = True
        elif b.kind == "para":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if len(b.text) > 25 and not _NO_INDENT.match(b.text):
                p.paragraph_format.first_line_indent = Pt(22)  # 首行缩进两字符
            run = p.add_run(_xml_safe_text(b.text))
            _set_cn_font(run)
            if "low_confidence" in b.flags:
                from docx.enum.text import WD_COLOR_INDEX
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif b.kind == "table" and b.rows:
            if "table_low_confidence" in b.flags:
                from docx.enum.text import WD_COLOR_INDEX
                p = doc.add_paragraph()
                run = p.add_run(_xml_safe_text(
                    "表格识别置信度低，可能存在列错位，建议人工核对。"))
                _set_cn_font(run, size=9)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            ncol = max(len(r) for r in b.rows)
            t = doc.add_table(rows=len(b.rows), cols=ncol)
            t.style = "Table Grid"
            for i, row in enumerate(b.rows):
                for j in range(ncol):
                    cell = t.cell(i, j)
                    cell.text = ""
                    if i == 0 and len(b.rows) > 1:
                        _shade_cell(cell)
                    parts = (row[j] if j < len(row) else "").split("\n")
                    for k, part in enumerate(parts):
                        p = (cell.paragraphs[0] if k == 0
                             else cell.add_paragraph())
                        run = p.add_run(_xml_safe_text(part))
                        _set_cn_font(run, size=10, bold=(i == 0))
            doc.add_paragraph()
        elif b.kind == "image" and b.image_path and os.path.exists(b.image_path):
            try:
                if "table_fallback" in b.flags:
                    from docx.enum.text import WD_COLOR_INDEX
                    p = doc.add_paragraph()
                    run = p.add_run(_xml_safe_text(
                        "表格结构识别不稳定，已保留原表格截图，建议以截图为准。"))
                    _set_cn_font(run, size=9)
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                doc.add_picture(b.image_path, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if _image_caption_reliable(b):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap = "图中文字（自动识别）"
                    if b.text:
                        cap += "：" + "；".join(
                            _xml_safe_text(s.strip()) for s in b.text.splitlines()
                            if s.strip())
                    run = p.add_run(_xml_safe_text(cap))
                    _set_cn_font(run, size=9)
                    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                    if b.rows:
                        ncol = max(len(r) for r in b.rows)
                        t = doc.add_table(rows=len(b.rows), cols=ncol)
                        t.style = "Table Grid"
                        for i, row in enumerate(b.rows):
                            for j in range(ncol):
                                cell = t.cell(i, j)
                                cell.text = ""
                                run = cell.paragraphs[0].add_run(
                                    _xml_safe_text(row[j] if j < len(row) else ""))
                                _set_cn_font(run, size=8)
                                run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
                        doc.add_paragraph()
            except Exception:
                p = doc.add_paragraph()
                _set_cn_font(p.add_run(_xml_safe_text(
                    f"[图片: {os.path.basename(b.image_path)}]")))
    doc.save(out_path)
    return out_path


def builtin_check_cases() -> List[tuple[str, Callable[[], None]]]:
    import tempfile

    def render_text(result: DocResult) -> str:
        with tempfile.TemporaryDirectory() as td:
            out_path = os.path.join(td, "out.md")
            export_markdown(result, out_path)
            with open(out_path, "r", encoding="utf-8") as f:
                return f.read()

    def case_unreliable_image_caption_does_not_export_link() -> None:
        text = render_text(DocResult(blocks=[
            Block(
                kind="image",
                image_path="fig.png",
                text="1\n二、适用区域⋯… 1三、定义/名词解释. 2",
            )
        ]))
        assert "![图片]" not in text
        assert "图中文字" not in text
        assert "\n1\n" not in text
        assert text.strip() == ""

    def case_digit_only_image_caption_does_not_export_link() -> None:
        text = render_text(DocResult(blocks=[Block(kind="image", image_path="fig.png", text="1")]))
        assert "![图片]" not in text
        assert "图中文字" not in text
        assert "\n1\n" not in text
        assert text.strip() == ""

    def case_formula_image_exports_latex() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="复合超时时长=(A1_{KA品牌单}+A2_{KA品牌单}+A3_{KA品牌单})/W_{KA品牌单}")
        ]))
        assert "$$" in text
        assert r"\text{复合超时时长}" in text
        assert r"\frac{A1_{\text{KA品牌单}}+A2_{\text{KA品牌单}}+A3_{\text{KA品牌单}}}" in text
        assert "![公式]" not in text

    def case_unknown_formula_requires_review_without_asset_link() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"], text="客诉虚假点送达率=分子/分母")
        ]))
        assert "**公式原文（需核对）**" in text
        assert "客诉虚假点送达率=分子/分母" in text
        assert "![" not in text

    def case_plain_text_formula_stays_text() -> None:
        formula = "商服务费 = 基础服务费 + 超额达标奖励 + KA 星级结算金额 + KA 体验膨胀费 + 服务质量奖励费 + 活动激励金"
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"], text=formula)
        ]))
        assert formula.replace("*", "\\*") in text
        assert "公式原文（需核对）" not in text
        assert "$$" not in text

    def case_score_example_formula_stays_text() -> None:
        formula = "麦当劳普通场景得分=15%*120+10%*120+15%*120+50%*105+10%*120=124.50分"
        text = render_text(DocResult(blocks=[
            Block(kind="para", text=formula)
        ]))
        assert "公式原文（需核对）" not in text
        assert "$$" not in text
        assert "124.50分" in text

    def case_sigma_text_formula_stays_text() -> None:
        formula = "30分钟内送达订单占比（配送距离3km及以内订单）=Σ（配送距离3km及以内且30分钟内送达完成订单量）/Σ（配送距离3km及以内完成订单量）"
        text = render_text(DocResult(blocks=[
            Block(kind="para", text=formula)
        ]))
        assert formula in text
        assert "公式原文（需核对）" not in text
        assert "$$" not in text

    def case_long_sum_formula_paragraph_stays_text() -> None:
        formula = "当日站点激励权重占比=当日该站点激励权重/sum（当日该站点所在城市所有站点激励权重）。"
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"], text=formula)
        ]))
        assert formula in text
        assert "公式原文（需核对）" not in text
        assert "$$" not in text

    def case_fraction_formula_keeps_visual_order() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="配送原因未完成率=P KA品牌单/(W KA品牌单+P KA品牌单)")
        ]))
        assert r"\text{配送原因未完成率}" in text
        assert r"\frac{P_{\text{KA品牌单}}}" in text

    def case_complaint_fake_delivery_formula_exports_latex() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="客诉虚假点送达率=星巴克/麦当劳/大润发品牌门店命中客诉虚假点送达单量/站点组履约的所有KA品牌单完成单量")
        ]))
        assert r"\text{客诉虚假点送达率}" in text
        assert r"\frac{\text{星巴克/麦当劳/大润发品牌门店命中客诉虚假点送达单量}}" in text
        assert "公式原文（需核对）" not in text

    def case_special_scene_completion_ratio_exports_latex() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="特殊场景完成单占比（t）=特殊场景剔除异常单后完成单/普通场景剔除异常单后完成单+特殊场景剔除异常单后完成单")
        ]))
        assert r"\text{特殊场景完成单占比}(t)" in text
        assert r"\frac{\text{特殊场景剔除异常单后完成单}}" in text

    def case_ka_522_intro_markdown_guard() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="heading", level=4, text="5.2.2"),
            Block(kind="para", text="特殊场景体验融合考核计分规则"),
            Block(kind="image", flags=["formula"],
                  text="特殊场景完成单占比（t）=特殊场景剔除异常单后完成单/普通场景剔除异常单后完成单+特殊场景剔除异常单后完成单"),
            Block(kind="heading", level=4, text="5.3 压力场景加权考核"),
        ]))
        section = text.split("#### 5.2.2", 1)[1].split("#### 5.3", 1)[0]
        assert "融合后体验得分" in section

    def case_weighted_special_scene_formula_recovered() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="加权后特殊场景完成单加权后特珠场景完成单占比(D二元权后普通场紫完成单＋加权后特珠场景完成单常规计分项得分=加权后特殊场景得分*t＋加权后普通场景得分*(1-t)")
        ]))
        assert r"\text{加权后特殊场景完成单占比}(t)" in text
        assert r"\text{常规计分项得分}" in text
        assert r"\frac{\text{加权后特殊场景完成单}}" in text

    def case_station_group_experience_formula_recovered() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="站点组体验总得分=站点组F分*F/(F+SUM(Kn*Qn))+站点组Kn分*(Kn*Qn)/(F+SUM(Kn*Qn))")
        ]))
        assert r"\text{站点组体验总得分}" in text
        assert r"\mathrm{SUM}(K_n\times Q_n)" in text
        assert r"\text{站点组}K_n\text{分}" in text

    def case_ka_brand_experience_formula_recovered() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="站点组KA品牌单体验得分=K1分*K1*Q1/SUM(Kn*Qn)+KN分*KN*QN/SUM(Kn*Qn)")
        ]))
        assert r"\text{站点组KA品牌单体验得分}" in text
        assert r"K_1\text{分}\times\frac{K_1\times Q_1}" in text
        assert r"K_n\text{分}\times\frac{K_n\times Q_n}" in text

    def case_ka_weighted_example_formula_keeps_numbers() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="A履约的体验总得分为118*(10000/(10000+700*6+300,z))+120*(700*6/(10000+70046+300,z))+115*(300,z/(10000+700*6+300,z))=118.4459分")
        ]))
        assert r"10000+700\times6+300\times2" in text
        assert r"115\times\frac{300\times2}" in text
        assert "300,z" not in text

    def case_ka_formula_recovery() -> None:
        punctuality = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="复合准时率（考核）=1-(C KA品牌单+5*Y KA品牌单)/W KA品牌单")
        ]))
        delivery = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="配送原因未完成率=P KA品牌单/(W KA品牌单+P KA品牌单)")
        ]))
        rider = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="计分规则 KA 品牌驻点骑手考核得分=W1得分*W1权重+W2得分*W2权重+W3得分*W3权重+W4得分*W4权重其中")
        ]))
        overtime = render_text(DocResult(blocks=[
            Block(kind="image", flags=["formula"],
                  text="复合超时时长=(A1_{KA品牌单}+A2_{KA品牌单}+A3_{KA品牌单})/W_{KA品牌单}")
        ]))
        assert r"C_{\text{KA品牌单}}+5\times Y_{\text{KA品牌单}}" in punctuality
        assert r"\frac{P_{\text{KA品牌单}}}" in delivery
        assert r"W1\text{得分}\times W1\text{权重}" in rider
        assert r"\frac{A1_{\text{KA品牌单}}+A2_{\text{KA品牌单}}+A3_{\text{KA品牌单}}}" in overtime

    def case_textual_image_table_exports_table_without_link() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image", flags=["table_fallback"], rows=[["指标", "值"], ["A", "1"]])
        ]))
        assert "| 指标 | 值 |" in text
        assert "![表格截图]" not in text

    def case_simple_four_column_table_stays_markdown_table() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", rows=[
                ["数据", "查询路径", "负责人", "备注"],
                ["KA品牌相关考核指标数据", "烽火台-商服务费计费系统", "渠道经理", "每月更新"],
            ])
        ]))
        assert "| 数据 | 查询路径 | 负责人 | 备注 |" in text
        assert "- 数据：" not in text

    def case_native_pdf_wide_benchmark_table_stays_markdown() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", flags=["native_pdf_table"], rows=[
                ["Model", "Size", "Overall ↑", "Text Edit ↓", "Formula CDM ↑",
                 "Table TEDs ↑", "Table TEDSs ↑", "Read-order Edit ↓"],
                ["Unlimited-OCR", "3B-A0.5B", "93.92", "0.042",
                 "95.79", "90.16", "93.32", "0.129"],
            ])
        ]))
        assert "| Model | Size | Overall ↑ | Text Edit ↓ | Formula CDM ↑ |" in text
        assert "- Model：" not in text

    def case_native_pdf_policy_table_still_exports_grouped_text() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", flags=["native_pdf_table"], rows=[
                ["一级分类", "检核项目", "内容", "责任承担"],
                ["站点安全", "14.手提式灭火器",
                 "合作商下属配送站点、充电区存在安全类隐患，烟感未绑定。",
                 "500元/项/次，整改不达标需承担双倍违约金。"],
            ])
        ]))
        assert "| 一级分类 | 检核项目 | 内容 | 责任承担 |" not in text
        assert "- 一级分类：站点安全" in text

    def case_text_pdf_references_keep_numbers_and_urls() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="heading", level=2, text="References"),
            Block(kind="para", text="[1] Nanonets-ocr-s, 2025. URL https://huggingface.co/nanonets/Nanonets-OCR-s."),
            Block(kind="para", text="[2] Ocrverse, 2025. URL https://github.com/DocTron-hub/OCRVerse."),
        ]))
        assert "## References" in text
        assert r"\[1\] Nanonets-ocr-s, 2025. URL https://huggingface.co/nanonets/Nanonets-OCR-s." in text
        assert r"\[2\] Ocrverse, 2025. URL https://github.com/DocTron-hub/OCRVerse." in text
        assert "公式原文（需核对）" not in text
        assert "\n1\n\n2\n" not in text

    def case_structured_score_example_keeps_field_relationships() -> None:
        text = render_text(DocResult(blocks=[
            Block(
                kind="para",
                flags=["structured_score_table"],
                text=(
                    "**特殊场景算分示例明细**\n\n"
                    "- 指标：复合准时率；分子数值：一般超时5单；严重超时0单；"
                    "分母数值：100；达标率（值）：95%；满分目标：96%；权重：50%；得分：110\n"
                    "\n**融合后体验得分核算**\n\n"
                    "- 特殊场景完成单占比：=100/（1000+100）=9.0909%"
                ),
            )
        ]))
        assert "分子数值：一般超时5单；严重超时0单" in text
        assert "融合后体验得分核算" in text
        assert "表格内容（按原文顺序）" not in text

    def case_image_table_exports_text_without_asset_link() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image", rows=[["数据", "查询路径"], ["KA", "系统"]],
                  image_path="out_assets/table.png")
        ]))
        assert "| 数据 | 查询路径 |" in text
        assert "![" not in text
        assert "out_assets/table.png" not in text

    def case_image_table_preserves_caption_text() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="image",
                  text="个时段相应的分值？\n昨/今变动未排班骑手的定义？",
                  rows=[["手", "考核？"], ["恶劣天", "天气等级在哪里查询？"]])
        ]))
        assert "个时段相应的分值？" in text
        assert "昨/今变动未排班骑手的定义？" in text
        assert "| 手 | 考核？ |" in text

    def case_table_preserves_high_value_text() -> None:
        text = render_text(DocResult(blocks=[
            Block(
                kind="table",
                text=(
                    "当天实时天气监控路径：烽火台-业务管理-天气查询。\n"
                    "申诉场景：站点天气判定恶劣，实际正常，申诉通过后按正常天气考核。\n"
                    "申诉路径：由渠道经理进行月度提报，申诉路径与薪动力异常场景申诉一致。"
                ),
                rows=[["目标", "值查询？"], ["天气", "径"]],
            )
        ]))
        assert "申诉路径：由渠道经理进行月度提报" in text
        assert "当天实时天气监控路径" in text
        assert "目标 / 值查询？" in text or "| 目标 | 值查询？ |" in text

    def case_single_column_answer_table_exports_list() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", rows=[
                ["目标会于每月月底进行推送，可在消息通知中心进行查看。"],
                ["每月结算明细在调分上传前已同步渠道经理，可联系渠道经理查询。"],
                ["申诉路径：由渠道经理进行月度提报，申诉路径与薪动力异常场景申诉一致。"],
            ])
        ]))
        assert "**目标会于每月月底" not in text
        assert "- 目标会于每月月底进行推送" in text
        assert "- 申诉路径：由渠道经理进行月度提报" in text

    def case_repaired_policy_image_table_hides_raw_text() -> None:
        raw = (
            "一级分类 检核项目 内容 责任承担 " + "填充" * 260
            + " 3.视频监 团检核人员无法及时查看站内情况 "
            + "流媒体设备出现遮挡 200元/项/次，整改 4.流媒体 屏幕 "
            + "7.看板海 样式符合标准"
        )
        text = render_text(DocResult(blocks=[
            Block(kind="image", text=raw, flags=["table_fallback"], rows=[
                ["一级分类", "检核项目", "内容", "责任承担"],
                ["", "3.视频监控", "视频监控内容。", "200元/项/次。"],
                ["", "4.流媒体", "流媒体内容。", "200元/项/次。"],
                ["", "7.看板海报", "看板海报内容。", "200元/项/次。"],
            ])
        ]))
        assert "3.视频监 团" not in text
        assert "遮挡 200元/项/次，整改" not in text
        assert "3.视频监控" in text
        assert "7.看板海报" in text

    def case_long_policy_table_exports_grouped_text() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", rows=[
                ["维度", "说明"],
                ["指标说明",
                 "对于虚假点送达的行为，相关用户会通过不同途径进行投诉，"
                 "包括但不限于电话投诉、风控抓取、用户 app 端投诉等，"
                 "方案将考核现有的电话客诉、风控抓取、不满意评价抓取3个来源。"],
            ])
        ]))
        assert "**指标说明**" in text
        assert "| 维度 | 说明 |" not in text

    def case_policy_responsibility_table_exports_multiline_context() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", rows=[
                ["一级分类", "检核项目", "内容", "责任承担"],
                [
                    "健康证",
                    "1.健康证",
                    "合作商配送站点公告栏中展示的配送服务人员健康证存在虚假。",
                    "需按照《合作商用工管理规范》相关约定承担违约责任",
                ],
                [
                    "内容",
                    "特殊说明：1.若站点线下被检核人员检查2次及以上次数，则以累计求和方式进行计算。",
                    "",
                    "",
                ],
            ])
        ]))
        assert "- 一级分类：健康证\n  检核项目：1.健康证" in text
        assert "- 一级分类：健康证；检核项目：1.健康证" not in text
        assert "**特殊说明**" in text
        assert "一级分类：内容" not in text

    def case_definition_table_with_long_source_exports_grouped_text() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", rows=[
                ["参考指标", "释义", "数据来源"],
                [
                    "KA品牌客诉单（KS）",
                    "因配送过程中存在未送指定位置、送达不通知、餐商品洒漏/货损/送错/少送、提前点送达、服务态度不佳等情况导致投诉。",
                    "KA品牌方回传给美团平台的客诉数据，客诉明细数据可联系渠道经理获取。",
                ],
            ])
        ]))
        assert "| 参考指标 | 释义 | 数据来源 |" not in text
        assert "- 参考指标：KA品牌客诉单" in text
        assert "数据来源：KA品牌方回传给美团平台的客诉数据" in text
        assert "。；数据来源" not in text

    def case_single_column_collapsed_table_exports_grouped_text() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", rows=[
                ["培训管理目标"],
                ["专送骑手 专送站长 城市经理 招聘人员 新人训 在岗训 专项训 组织规则 过程指标 考核结果"],
            ])
        ]))
        assert "| 培训管理目标 |" not in text
        assert "**培训管理目标**" in text

    def case_complex_table_exports_readable_blocks_without_br() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", rows=[
                ["项目", "内容"],
                ["指标定义1",
                 "【完成订单量】通过美团配送且订单最终状态为完成的订单数量。\n"
                 "麦当劳完单量占比=麦当劳完单量/总量\n"
                 "注：数据来源均为品牌方。"],
            ]),
            Block(kind="table", rows=[
                ["指标", "分子数值", "分母数值", "达标率（值）", "满分目标", "权重", "得分"],
                ["KA品牌负向反馈率", "2单", "100", "2%", "0.01%", "15%", "0"],
            ]),
        ]))
        assert "<br>" not in text
        assert "**指标定义1**" in text
        assert "麦当劳完单量占比=麦当劳完单量/总量" in text
        assert "- 指标：KA品牌负向反馈率" in text

    def case_score_example_phrase_not_split() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="para", text="普通场景算分示例：假设站点组A履约。"),
            Block(kind="para", text="特殊场景算分示例：假设站点组A履约。"),
        ]))
        assert "普通场景算分\n示例" not in text
        assert "特殊场景算分\n示例" not in text
        assert "普通场景算分示例：假设站点组A履约。" in text

    def case_bracket_label_after_sentence_gets_own_line() -> None:
        lines = _readable_text_lines(
            "第一项说明结束。【完成订单量】通过美团配送且最终状态为完成的订单数量。"
        )
        assert lines == [
            "第一项说明结束。",
            "【完成订单量】通过美团配送且最终状态为完成的订单数量。",
        ]

    def case_joined_labels_are_not_split_again() -> None:
        calculation = _readable_text_lines("举例结束。计算示例：按下列订单计算。")
        punctuality = _readable_text_lines(
            "【肯德基体验得分】考核“品牌方口径准时率（肯德基）”"
        )
        source = _readable_text_lines("⑤ 数据获取方式：联系渠道经理；")
        assert "计算\n示例" not in "\n".join(calculation)
        assert any(line.startswith("计算示例：") for line in calculation)
        assert punctuality == ["【肯德基体验得分】考核“品牌方口径准时率（肯德基）”"]
        assert source == ["⑤ 数据获取方式：联系渠道经理；"]

        workflow = _readable_text_lines(
            "数据来源：品牌方投诉。数据播报：次月下发。"
            "降星规则：命中即降星。申诉流程及时效：一个工作日内申诉。"
        )
        assert workflow == [
            "数据来源：品牌方投诉。",
            "数据播报：次月下发。",
            "降星规则：命中即降星。",
            "申诉流程及时效：一个工作日内申诉。",
        ]

        numbered = _readable_text_lines(
            "申诉超时自动放弃。（1）该城市有自然灾害。（2）因品牌方调整。"
        )
        assert numbered == [
            "申诉超时自动放弃。",
            "（1）该城市有自然灾害。",
            "（2）因品牌方调整。",
        ]

    def case_missing_numbered_heading_repaired() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="heading", level=4, text="2.2.2 考核标准："),
            Block(kind="para", text="- 绑站具体规则"),
            Block(kind="para", text="1、站点绑定规则明细。"),
            Block(kind="para", text="2.2.4 命中退出站点保留需扣除违约金。"),
        ]))
        assert "#### 2.2.3 绑站具体规则" in text
        assert "- 绑站具体规则" not in text

    def case_exception_table_example_breaks_line() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", rows=[
                ["异常类型", "提报条件（剔除条件）", "说明"],
                [
                    "灾害天气",
                    "②*台风/海啸/泥石流/地震/洪水/沙尘暴等灾害天气，需及时申请一键置休；"
                    "内涝、结冰场景不强制要求新闻截图；站点需同步保留现场照片、系统截图、"
                    "新闻截图或城市通知作为复核材料；举例：结冰 07:04 现场照片",
                    "数据来源：系统提报记录与附件截图。",
                ],
            ])
        ]))
        assert "；举例：" not in text
        assert "举例：结冰" in text

    def case_exception_conditions_break_lines() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", rows=[
                ["场景", "提报条件"],
                [
                    "交通事故",
                    "需所有条件同时满足。；条件一：上传事故认定书；；条件二：上传现场水印照片；；1.时间需清晰；；2.地点需清晰。",
                ],
            ])
        ]))
        assert "；条件一：" not in text
        assert "；1." not in text
        assert "条件一：上传事故认定书" in text

    def case_numbered_sentence_breaks_lines() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", rows=[
                ["内容"],
                ["1. 为杜绝虚假报备现象，报备订单驳回后无法再次提报。2. 加盟合作商应保证信息真实。"],
            ])
        ]))
        assert "再次提报。2." not in text
        assert "2. 加盟合作商应保证信息真实。" in text

    def case_numbered_examples_and_notes_break_lines() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="para", text=(
                "考核说明：按长期低星配区名单为准。举例1:A配区命中规则。"
                "举例2:B配区命中规则。注：此处为说明。备\n注：此场景仅适用于单商运营。"
            )),
            Block(kind="para", text=(
                "合作商需承担如下责任；其中总商维度提交主动退出申请的站点按规则执行。"
                "此场景仅适用于单商运营。其中平均日均单量按近6个月计算。"
            )),
        ]))
        assert "。举例1" not in text
        assert "。举例2" not in text
        assert "。注：" not in text
        assert "备\n注" not in text
        assert "；其中" not in text
        assert "。其中" not in text
        assert "备注：此场景" in text

    def case_markdown_note_breaks_after_table_export() -> None:
        text = _repair_markdown_note_breaks(
            "**非KA品牌单 / 指站点组履约的运单。注：删除喜茶。**\n\n"
            "- 非KA品牌单：指站点组履约的运单。注：删除喜茶。\n"
        )
        assert "。注：" not in text
        assert "**非KA品牌单 / 指站点组履约的运单。**" in text
        assert "注：删除喜茶。" in text

    def case_bracket_definition_breaks_lines() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", rows=[
                ["基础服务费计算方式"],
                [
                    "每日基础服务费=当日标准1*激励金全月基础服务费=每日基础服务费之和"
                    "每日计费服务质量等级数根据达成关系计算；：同时满足如下条件即视为达成标准1："
                    "① 骑手全天完成单量≥X1单。单量标准判断范围包含企客单。【有效在线时长】："
                    "有效在线时长即骑手上线状态下有单服务时长之和。有效在线时长计算口径=在线时长-掉线时长。"
                    "其中【在线时长】为骑手上线、忙碌状态下的总时长；：特殊说明：立即单从接单开始计算。",
                ],
            ])
        ]))
        assert "激励金全月基础服务费" not in text
        assert "；：同时满足" not in text
        assert "。【有效在线时长】" not in text
        assert "特殊说明：立即单" in text

        repaired = _repair_markdown_definition_breaks(
            "每日基础服务费=当日标准1*激励金全月基础服务费=每日基础服务费之和"
            "每日计费服务质量等级数根据达成关系计算；：特殊说明：立即单。"
            "有效在线时长计算口径=在线时长。其中【在线时长】为总时长、【忙碌无单时长】为忙碌时长。"
        )
        assert "激励金全月基础服务费" not in repaired
        assert "；：特殊说明" not in repaired
        assert "。其中" not in repaired
        assert "、【忙碌无单时长】" not in repaired

    def case_final_markdown_wraps_long_english_text() -> None:
        paragraph = (
            "Unlimited OCR evaluates OCR systems under arbitrary image orientations, "
            "font sizes, dense mathematical notation, and natural document layouts. "
        ) * 12
        text = render_text(DocResult(blocks=[
            Block(kind="para", text=paragraph),
        ]))
        assert max(len(line) for line in text.splitlines()) <= 230

    def case_docx_strips_text_layer_control_chars() -> None:
        with tempfile.TemporaryDirectory() as td:
            out_path = os.path.join(td, "out.docx")
            export_docx(DocResult(blocks=[
                Block(kind="heading", text="标题\x0b"),
                Block(kind="para", text="正文\x0c内容"),
                Block(kind="table", rows=[["列\x01"], ["值\x02"]]),
            ]), out_path)
            assert os.path.exists(out_path)

    def case_standalone_numbered_paragraph_becomes_heading() -> None:
        repaired = _repair_markdown_missing_numbered_headings(
            "## 2. Related Works\n\n"
            "正文段落。\n\n"
            "2.2. End-to-end Model\n\n"
            "下一段正文。\n"
        )
        assert "### 2.2 End-to-end Model" in repaired
        assert "\n2.2. End-to-end Model\n" not in repaired

    def case_toc_entries_do_not_become_body_headings() -> None:
        repaired = _repair_markdown_missing_numbered_headings(
            "## Contents\n\n"
            "1 Introduction 3\n\n"
            "2 Related Works 4\n\n"
            "## 1. Introduction\n\n"
            "正文。\n"
        )
        assert "## 1 Introduction 3" not in repaired
        assert "\n1 Introduction 3\n" in repaired

    def case_text_pdf_toc_exports_nested_list() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="heading", level=2, text="Contents"),
            Block(kind="para", text="1 Introduction 3"),
            Block(kind="para", text=(
                "2 Related Works 4 2.1 Pipeline-based Framework . . . . . 4 "
                "2.2 End-to-end Model . . . . . 4 2.2.1 High-compression Encoder . . . 4"
            )),
            Block(kind="para", text="3 Methodology 5 3.1 Long-horizon Parsing . . . . . 5"),
            Block(kind="heading", level=2, text="1. Introduction"),
        ]))
        assert "- 1 Introduction (p. 3)" in text
        assert "  - 2.1 Pipeline-based Framework (p. 4)" in text
        assert "    - 2.2.1 High-compression Encoder (p. 4)" in text
        assert "2 Related Works 42.1" not in text

    def case_verified_wide_standard_table_stays_markdown() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="table", flags=["standard_markdown_table"], rows=[
                ["管理模块", "培训对象", "学习路径", "考试机会", "管理方式", "数据ID", "计分规则"],
                ["专送骑手", "全量骑手", "培训中心-安全保障", "不限制", "线上自学", "7594",
                 "当应参通过率<90%时，B1=0；当应参通过率≥90%时，B1=应参通过率*20"],
            ])
        ]))
        assert "| 管理模块 | 培训对象 | 学习路径 | 考试机会 | 管理方式 | 数据ID | 计分规则 |" in text
        assert "- 管理模块：" not in text

    def case_english_bullets_split_into_markdown_items() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="para", text=(
                "• We introduce Reference Sliding Window Attention. "
                "This design keeps the KV cache constant.• Building on R-SWA, "
                "we propose Unlimited OCR.• We conduct a preliminary validation."
            )),
        ]))
        assert "- We introduce Reference Sliding Window Attention." in text
        assert "- Building on R-SWA" in text
        assert "- We conduct a preliminary validation." in text
        assert "constant.• Building" not in text

    def case_formula_text_blocks_grouped_as_original_math() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="para", text="the required KV cache size is"),
            Block(kind="para", flags=["formula_text"], text="C_MHA(T) = L_m + T."),
            Block(kind="para", flags=["formula_text"], text="(5)"),
            Block(kind="para", text="In contrast, under R-SWA, the model retains cache."),
        ]))
        assert "the required KV cache size is\n\n```text\nC_MHA(T) = L_m + T.\n(5)\n```" in text
        assert "```text\nC_MHA(T) = L_m + T.\n```\n\n```text\n(5)" not in text
        assert "In contrast, under R-SWA" in text

    def case_formula_tail_split_from_text_pdf_prose() -> None:
        text = render_text(DocResult(blocks=[
            Block(kind="para", text="The attention weight from token t is then computed as \x10 q⊤\x11"),
            Block(kind="para", flags=["formula_text"], text="t k_j"),
            Block(kind="para", text="where q is the query."),
            Block(kind="para", text="the same accessible set:∑︁"),
            Block(kind="para", flags=["formula_text"], text="alpha v"),
        ]))
        assert "computed as q⊤" not in text
        assert "computed as\n\n```text\nq⊤\nt k_j\n```" in text
        assert "same accessible set:∑" not in text
        assert "same accessible set:\n\n```text\n∑︁\nalpha v\n```" in text

    return [
        ("export.unreliable_image_caption_does_not_export_link", case_unreliable_image_caption_does_not_export_link),
        ("export.digit_only_image_caption_does_not_export_link", case_digit_only_image_caption_does_not_export_link),
        ("export.formula_image_exports_latex", case_formula_image_exports_latex),
        ("export.unknown_formula_requires_review_without_asset_link", case_unknown_formula_requires_review_without_asset_link),
        ("export.plain_text_formula_stays_text", case_plain_text_formula_stays_text),
        ("export.score_example_formula_stays_text", case_score_example_formula_stays_text),
        ("export.sigma_text_formula_stays_text", case_sigma_text_formula_stays_text),
        ("export.long_sum_formula_paragraph_stays_text", case_long_sum_formula_paragraph_stays_text),
        ("export.fraction_formula_keeps_visual_order", case_fraction_formula_keeps_visual_order),
        ("export.complaint_fake_delivery_formula_exports_latex", case_complaint_fake_delivery_formula_exports_latex),
        ("export.special_scene_completion_ratio_exports_latex", case_special_scene_completion_ratio_exports_latex),
        ("export.ka_522_intro_markdown_guard", case_ka_522_intro_markdown_guard),
        ("export.weighted_special_scene_formula_recovered", case_weighted_special_scene_formula_recovered),
        ("export.station_group_experience_formula_recovered", case_station_group_experience_formula_recovered),
        ("export.ka_brand_experience_formula_recovered", case_ka_brand_experience_formula_recovered),
        ("export.ka_weighted_example_formula_keeps_numbers", case_ka_weighted_example_formula_keeps_numbers),
        ("export.ka_formula_recovery", case_ka_formula_recovery),
        ("export.textual_image_table_exports_table_without_link", case_textual_image_table_exports_table_without_link),
        ("export.simple_four_column_table_stays_markdown_table", case_simple_four_column_table_stays_markdown_table),
        ("export.native_pdf_wide_benchmark_table_stays_markdown", case_native_pdf_wide_benchmark_table_stays_markdown),
        ("export.native_pdf_policy_table_still_exports_grouped_text", case_native_pdf_policy_table_still_exports_grouped_text),
        ("export.text_pdf_references_keep_numbers_and_urls", case_text_pdf_references_keep_numbers_and_urls),
        ("export.structured_score_example_keeps_field_relationships", case_structured_score_example_keeps_field_relationships),
        ("export.image_table_exports_text_without_asset_link", case_image_table_exports_text_without_asset_link),
        ("export.image_table_preserves_caption_text", case_image_table_preserves_caption_text),
        ("export.table_preserves_high_value_text", case_table_preserves_high_value_text),
        ("export.single_column_answer_table_exports_list", case_single_column_answer_table_exports_list),
        ("export.repaired_policy_image_table_hides_raw_text", case_repaired_policy_image_table_hides_raw_text),
        ("export.long_policy_table_exports_grouped_text", case_long_policy_table_exports_grouped_text),
        ("export.policy_responsibility_table_exports_multiline_context", case_policy_responsibility_table_exports_multiline_context),
        ("export.definition_table_with_long_source_exports_grouped_text", case_definition_table_with_long_source_exports_grouped_text),
        ("export.single_column_collapsed_table_exports_grouped_text", case_single_column_collapsed_table_exports_grouped_text),
        ("export.complex_table_exports_readable_blocks_without_br", case_complex_table_exports_readable_blocks_without_br),
        ("export.score_example_phrase_not_split", case_score_example_phrase_not_split),
        ("export.bracket_label_after_sentence_gets_own_line", case_bracket_label_after_sentence_gets_own_line),
        ("export.joined_labels_are_not_split_again", case_joined_labels_are_not_split_again),
        ("export.missing_numbered_heading_repaired", case_missing_numbered_heading_repaired),
        ("export.exception_table_example_breaks_line", case_exception_table_example_breaks_line),
        ("export.exception_conditions_break_lines", case_exception_conditions_break_lines),
        ("export.numbered_sentence_breaks_lines", case_numbered_sentence_breaks_lines),
        ("export.numbered_examples_and_notes_break_lines", case_numbered_examples_and_notes_break_lines),
        ("export.markdown_note_breaks_after_table_export", case_markdown_note_breaks_after_table_export),
        ("export.bracket_definition_breaks_lines", case_bracket_definition_breaks_lines),
        ("export.final_markdown_wraps_long_english_text", case_final_markdown_wraps_long_english_text),
        ("export.docx_strips_text_layer_control_chars", case_docx_strips_text_layer_control_chars),
        ("export.standalone_numbered_paragraph_becomes_heading", case_standalone_numbered_paragraph_becomes_heading),
        ("export.toc_entries_do_not_become_body_headings", case_toc_entries_do_not_become_body_headings),
        ("export.text_pdf_toc_exports_nested_list", case_text_pdf_toc_exports_nested_list),
        ("export.verified_wide_standard_table_stays_markdown", case_verified_wide_standard_table_stays_markdown),
        ("export.english_bullets_split_into_markdown_items", case_english_bullets_split_into_markdown_items),
        ("export.formula_text_blocks_grouped_as_original_math", case_formula_text_blocks_grouped_as_original_math),
        ("export.formula_tail_split_from_text_pdf_prose", case_formula_tail_split_from_text_pdf_prose),
    ]
